# Imports 
from flask import Flask, render_template, request, redirect, url_for, flash, send_file, session
from flask_login import LoginManager, UserMixin, login_user, login_required, current_user, logout_user
from flask_mail import Mail, Message
from flask_pymongo import PyMongo
from docx import Document
import io
from flask_login import login_required
from werkzeug.security import generate_password_hash, check_password_hash
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
import io
from flask import send_file
from datetime import datetime  # Replace 'import datetime' with this
# app instance creation

from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
app = Flask(__name__)
# Secret key for session management
app.secret_key = 'your_secret_key'

# Initialize Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)

# User model (adjust it to your needs)
class User(UserMixin):
    def __init__(self, id, username):
        self.id = id
        self.username = username

# User loader callback function for Flask-Login
@login_manager.user_loader
def load_user(user_id):
    # In a real app, fetch the user from the database using user_id
    return User(user_id, "test_username")  # Replace with actual user retrieval logic

# Simulate a database query to get answers for the logged-in user
def get_answers_from_database():
    # Fetch answers based on the current user (you need to adjust this to your database logic)
    if current_user.is_authenticated:
        answers = mongo.db.answers.find({"user_id": current_user.id})  # Replace with your actual DB query
        return answers
    else:
        return None

# Gmail setup
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 465
app.config['MAIL_USE_SSL'] = True
app.config['MAIL_USERNAME'] = "vedant.22220153@viit.ac.in"
app.config['MAIL_PASSWORD'] = "bogt bjmz zayb xpyi"
app.config['MAIL_DEFAULT_SENDER'] = 'vedant.22220153@viit.ac.in'
mail = Mail(app)

# Mongo-DB compass setup
app.config["MONGO_URI"] = "mongodb://localhost:27017/Login"
app.secret_key = 'abcdefgh'
mongo = PyMongo(app)

# Homepage
@app.route("/")
def home():
    return render_template("login.html")

# SignUp route
@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")
        email = request.form.get("email")

        # Check if username already exists
        existing_user = mongo.db.accounts.find_one({"username": username})
        if existing_user:
            flash("Account already exists")
            return redirect(url_for('signup'))

        # Check if email already exists
        existing_mail = mongo.db.accounts.find_one({"email": email})
        if existing_mail:
            flash("Email already exists")
            return redirect(url_for('signup'))

        hashed_password = generate_password_hash(password)

        # Create a new user
        mongo.db.accounts.insert_one({
            "username": username,
            "password": hashed_password,
            "email": email})
        flash("Account created successfully")
        return redirect(url_for('login'))
    return render_template("signup.html")

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        # Fetch the user from the database
        user = mongo.db.accounts.find_one({"username": username})
        
        if user:
            # Verify the password
            if check_password_hash(user['password'], password):
                # Create a User object and log in
                user_obj = User(str(user['_id']), username)
                login_user(user_obj)
                return redirect(url_for('after_login'))  # Redirect to after_login route
            else:
                flash('Invalid password')  # Incorrect password
        else:
            flash('User not found')  # Username doesn't exist

    return render_template('login.html')

# @app.route('/after_login', methods=['GET', 'POST'])
# @login_required
# def after_login():
#     if request.method == 'POST':
#         # Collect form data from after-login.html
#         rfx_number = request.form.get('rfx')
#         etender_number = request.form.get('etender')
#         tendercall_number = request.form.get('tendercall')

#         # Update the existing doc in "answers" where user_id + tender=Proqurement
#         # so that we do NOT create a separate doc or collection
#         mongo.db.answers.update_one(
#             {
#                 "user_id": current_user.id,
#                 "tender": "Proqurement"
#             },
#             {
#                 "$set": {
#                     "rfx_number": rfx_number,
#                     "etender_number": etender_number,
#                     "tendercall_number": tendercall_number
#                 }
#             },
#             upsert=False  # upsert=False ensures we don't create a new doc if none is found
#         )

#         # Redirect or render next page after storing
#         return render_template('tender1.html')

#     # Render after-login.html for GET request
#     return render_template('after-login.html')

@app.route('/after_login', methods=['GET', 'POST'])
@login_required
def after_login():
    if request.method == 'POST':
        # Collect form data from after-login.html
        rfx_number = request.form.get('rfx')
        etender_number = request.form.get('etender')
        tendercall_number = request.form.get('tendercall')

        # Store submission details in the answers collection under "Proqurement"
        mongo.db.answers.update_one(
            {
                "user_id": current_user.id,
                "tender": "Proqurement"
            },
            {
                "$set": {
                    "submission_details.rfx_number": rfx_number,
                    "submission_details.etender_number": etender_number,
                    "submission_details.tendercall_number": tendercall_number,
                    "submitted_at": datetime.now()
                }
            },
            upsert=True  # Create if not exists
        )

        return render_template('tender1.html')

    return render_template('after-login.html')

@app.route('/logout')
def logout():
    logout_user()  # Flask-Login handles the logout
    return redirect(url_for('login'))

@app.route("/display_sub_form/<username>")
def display_sub_form(username):
    return render_template("display_sub_form.html", username=username)

@app.route("/submit", methods=["POST"])
@login_required
def tender():
    selected_sub = request.form.get("tender")
    print({selected_sub})
    answers = {
        "Question 1": request.form.get("q1"),
    }

    if selected_sub: 
        existing_submission = mongo.db.answers.find_one({"user_id": current_user.id, "tender": selected_sub})

        if existing_submission:
            # Update existing record
            mongo.db.answers.update_one(
                {"user_id": current_user.id, "tender": selected_sub},
                {"$set": {"answers": answers}}
            )
            flash("Your answers have been updated successfully!")
        else:
            # Insert new answers
            mongo.db.answers.insert_one({
                "user_id": current_user.id,
                "answers": answers,
                "tender": selected_sub,
            })
            flash("Answers submitted successfully!")

    # Redirect to the appropriate form based on tender
    if selected_sub == "Proqurement":
        return render_template("tender1.html")
    elif selected_sub == "Services":
        return render_template("tender2.html")
    elif selected_sub == "tender-3":
        return render_template("tender3.html")

    return redirect(url_for('display_sub_form', username=current_user.username))



@app.route('/tender1_form', methods=['POST'])
@login_required
def tender1_form():
    # Fetch the submitted form data
    answers = {
        "Question 1": request.form.get("q1"),
    }
    selected_sub = "Proqurement"  # Static identifier for this form

    existing_submission = mongo.db.answers.find_one({"user_id": current_user.id, "tender": selected_sub})

    if existing_submission:
        # Update existing answers
        mongo.db.answers.update_one(
            {"user_id": current_user.id, "tender": selected_sub},
            {"$set": {"answers": answers}}
        )
        flash("Your answers have been updated successfully!")
    else:
        # Insert new answers
        mongo.db.answers.insert_one({
            "user_id": current_user.id,
            "tender": selected_sub,
            # "answers": answers,
        })
        flash("Answers submitted successfully!")

    # Redirect to the next tender page
    return render_template('tender1_next.html')



@app.route('/tender1_next', methods=['GET', 'POST'])
@login_required
def tender1_next():
    if request.method == "POST":
        # 1. Collect the selected products
        selected_products = request.form.getlist("product")
        user_selections = []

        for prod in selected_products:
            # Now "prod" has 2 parts: e.g. "Material" and "ProductID"
            # e.g. "Wood-WD123" -> ["Wood", "WD123"]
            parts = prod.split("-")
            material = parts[0].strip()
            product_id = parts[1].strip()

            # Build field names for price & quantity
            price_field = f"price_{product_id}"
            quantity_field = f"quantity_{product_id}"

            # Parse the inputs from the form
            price_input = request.form.get(price_field, "0")
            qty_advertised_input = request.form.get(quantity_field, "0")

            user_selections.append({
                "material": material,
                "product_id": product_id,
                "price": price_input,
                "qty_advertised": qty_advertised_input
            })

        mongo.db.answers.update_one(
            {"user_id": current_user.id, "tender": "Proqurement"},
            {"$set": {"answers.selected_materials": user_selections}},
            upsert=True
        )

    return redirect(url_for('review'))


@app.route('/tender2_form', methods=['POST'])
@login_required
def tender2_form():
    answers = {
        "Question 1": request.form.get("q1"),
        "Question 2": request.form.get("q2"),
        "Question 3": request.form.get("q3"),
        "Question 4": request.form.get("q4"),
    }
    selected_sub = "tender-2"

    existing_submission = mongo.db.answers.find_one({"user_id": current_user.id, "tender": selected_sub})

    if existing_submission:
        mongo.db.answers.update_one(
            {"user_id": current_user.id, "tender": selected_sub},
            {"$set": {"answers": answers}}
        )
        flash("Your answers have been updated successfully!")
    else:
        mongo.db.answers.insert_one({
            "user_id": current_user.id,
            "tender": selected_sub,
            "answers": answers,
        })
        flash("Answers submitted successfully!")

    return render_template('tender2_next.html')

@app.route('/tender2_next', methods=['POST'])
@login_required
def tender2_next():
    answers = {
        "Question 1": request.form.get("q1"),
        "Question 2": request.form.get("q2"),
    }
    selected_sub = "tender-2-next"

    existing_submission = mongo.db.answers.find_one({"user_id": current_user.id, "tender": selected_sub})
    if existing_submission:
        mongo.db.answers.update_one(
            {"user_id": current_user.id, "tender": selected_sub},
            {"$set": {"answers": answers}}
        )
        flash("Your answers have been updated successfully!")
    else:
        mongo.db.answers.insert_one({
            "user_id": current_user.id,
            "tender": selected_sub,
            "answers": answers,
        })
        flash("Answers submitted successfully!")

    return redirect(url_for('review'))  # Redirect to tender 3 Next form


@app.route('/tender3_form', methods=['POST'])
@login_required
def tender3_form():
    answers = {
        "Question 1": request.form.get("q1"),
        "Question 2": request.form.get("q2"),
        "Question 3": request.form.get("q3"),
        "Question 4": request.form.get("q4"),
    }
    selected_sub = "tender-3"

    existing_submission = mongo.db.answers.find_one({"user_id": current_user.id, "tender": selected_sub})

    if existing_submission:
        mongo.db.answers.update_one(
            {"user_id": current_user.id, "tender": selected_sub},
            {"$set": {"answers": answers}}
        )
        flash("Your answers have been updated successfully!")
    else:
        mongo.db.answers.insert_one({
            "user_id": current_user.id,
            "tender": selected_sub,
            "answers": answers,
        })
        flash("Answers submitted successfully!")

    return render_template('tender3_next.html')  # Redirect to Review page

@app.route('/tender3_next', methods=['POST'])
@login_required
def tender3_next():
    answers = {
        "Question 1": request.form.get("q1"),
        "Question 2": request.form.get("q2"),
    }
    selected_sub = "tender-3-next"

    # if not all(answers.values()):
    #     flash("Please answer all questions before submitting.")
    #     return redirect(url_for('tender3_next'))

    # Check if the user already submitted answers
    existing_submission = mongo.db.answers.find_one({"user_id": current_user.id, "tender": selected_sub})
    if existing_submission:
        mongo.db.answers.update_one(
            {"user_id": current_user.id, "tender": selected_sub},
            {"$set": {"answers": answers}}
        )
        flash("Your answers have been updated successfully!")
    else:
        mongo.db.answers.insert_one({
            "user_id": current_user.id,
            "tender": selected_sub,
            "answers": answers,
        })
        flash("Answers submitted successfully!")

    return redirect(url_for('review'))  # Redirect to Review page


@app.route('/review', methods=['GET', 'POST'])
@login_required
def review():
    if request.method == 'POST':
        form_data = request.form.to_dict(flat=False)

        # Update submission details in answers collection
        if 'submissions[rfx_number]' in form_data:
            mongo.db.answers.update_one(
                {"user_id": current_user.id, "tender": "Proqurement"},
                {
                    "$set": {
                        "submission_details.rfx_number": form_data['submissions[rfx_number]'][0],
                        "submission_details.etender_number": form_data['submissions[etender_number]'][0],
                        "submission_details.tendercall_number": form_data['submissions[tendercall_number]'][0],
                        "submission_details.submitted_at": datetime.now()
                    }
                },
                upsert=True
            )

        # Update answers
        for tender_key in [key for key in form_data if key.startswith('grouped_answers[')]:
            tender = tender_key.split('[')[1].split(']')[0]
            answers_data = {}
            for key, value in form_data.items():
                if key.startswith(f'grouped_answers[{tender}]') and key != f'grouped_answers[{tender}][tender]':
                    question = key.split('[')[-1].rstrip(']')
                    answers_data[question] = value[0] if value else ""
            mongo.db.answers.update_one(
                {"user_id": current_user.id, "tender": tender},
                {"$set": {"answers": answers_data}},
                upsert=True
            )

        # Update selected materials
        selected_materials = []
        material_count = len(form_data.get('selected_materials[material][]', []))
        for i in range(material_count):
            material = {
                "material": form_data['selected_materials[material][]'][i],  # Fixed from rfx_number
                "product_id": form_data['selected_materials[product_id][]'][i],
                "qty_advertised": form_data['selected_materials[qty_advertised][]'][i],
                # "price" is not in the form, remove or adjust if needed
            }
            selected_materials.append(material)
        if selected_materials:
            mongo.db.answers.update_one(
                {"user_id": current_user.id, "tender": "Proqurement"},
                {"$set": {"answers.selected_materials": selected_materials}},
                upsert=True
            )

        flash("Answers and materials updated successfully!")
        return redirect(url_for('review'))

    # GET request - Fetch data from MongoDB
    answers_cursor = mongo.db.answers.find({"user_id": current_user.id})
    grouped_answers = {}
    selected_materials = []
    submissions = {}

    for ans_doc in answers_cursor:
        tender_name = ans_doc["tender"]
        if tender_name not in grouped_answers:
            grouped_answers[tender_name] = []
        grouped_answers[tender_name].append({
            "tender": tender_name,
            "answers": ans_doc.get("answers", {})
        })
        if tender_name == "Proqurement":
            if "submission_details" in ans_doc:
                submissions = ans_doc["submission_details"]
            if "selected_materials" in ans_doc.get("answers", {}):
                selected_materials = ans_doc["answers"]["selected_materials"]

    return render_template(
        "review.html",
        grouped_answers=grouped_answers,
        selected_materials=selected_materials,
        submissions=submissions
    )


@app.route('/update_answers', methods=['POST'])
@login_required
def update_answers():
    # Parse form data
    form_data = request.form.to_dict(flat=False)

    # Update answers for each tender
    for key, value in form_data.items():
        if key.startswith("tender["):  
            tender_index = key[8:-1]  
            tender_name = value[0]  

            # Collect answers for this tender
            updated_answers = {
                k.split("][", 1)[-1][:-1]: v[0]
                for k, v in form_data.items()
                if k.startswith(f"answers[{tender_index}]")
            }

            # Update the database for the current tender
            mongo.db.answers.update_one(
                {"user_id": current_user.id, "tender": tender_name},
                {"$set": {"answers": updated_answers}}
            )

    # Update selected materials in the database
    selected_materials = request.form.getlist("selected_materials")  # Get from form
    mongo.db.users.update_one(
        {"_id": current_user.id},
        {"$set": {"selected_materials": selected_materials}}
    )

    flash("Answers and materials updated successfully!")
    return redirect(url_for('review'))





@app.route('/generate_docx', methods=['GET'])
@login_required
def generate_docx():
    # Fetch answers and submissions for the current logged-in user
    answers = list(mongo.db.answers.find({"user_id": current_user.id}))
    submissions = mongo.db.submissions.find_one({"user_id": current_user.id})  # Assuming submissions are stored this way

    # Create a new Document
    doc = Document()
    doc.add_heading('Your Copy of Submitted Details', 0)

    # Submission Details Section
    if submissions:
        doc.add_heading('Submission Details', level=1)
        doc.add_paragraph(f"RFX Number: {submissions.get('rfx_number', 'N/A')}")
        doc.add_paragraph(f"E-Tender Number: {submissions.get('etender_number', 'N/A')}")
        doc.add_paragraph(f"Tender Call Number: {submissions.get('tendercall_number', 'N/A')}")
        doc.add_paragraph("")  # Add spacing

    # Answers Section
    if answers:
        for answer_set in answers:
            tender = answer_set.get("tender", "Unknown tender")
            doc.add_heading(f"Tender: {tender.capitalize()}", level=1)
            
            # Regular answers (excluding selected_materials)
            answers_dict = answer_set.get("answers", {})
            for question, answer in answers_dict.items():
                if question != "selected_materials":
                    doc.add_paragraph(f"{question}: {answer}")

            # Selected Materials Section
            selected_materials = answers_dict.get("selected_materials", [])
            if selected_materials:
                doc.add_heading("Selected Materials", level=2)
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Material Name'
                hdr_cells[1].text = 'Product ID'
                hdr_cells[2].text = 'Quantity'
                hdr_cells[3].text = 'Delivery Schedule'

                for item in selected_materials:
                    row_cells = table.add_row().cells
                    row_cells[0].text = item.get('material', '')
                    row_cells[1].text = item.get('product_id', '')
                    row_cells[2].text = item.get('qty_advertised', '')
                    row_cells[3].text = "CM: [    ]Nos. within 2 months from date of LOA\nCP: @ [   ] Nos. per month thereafter"
    else:
        doc.add_paragraph("No Detils available for review.")

    # Save the document to an in-memory file
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="review_answers.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Forgot-password route
@app.route("/forgot-password", methods=["GET", "POST"])
def forgot_password():
    if request.method == "POST":
        email = request.form.get("email")
        account = mongo.db.accounts.find_one({"email": email})

        if account:
            reset_link = url_for('reset_password', _external=True)
            
            # Create Message object
            msg = Message("Password Reset Request", sender="your-email@example.com", recipients=[email])
            msg.body = f'Hi, {account["username"]}!\n\nClick on the link to reset your password: \n{reset_link}'

            try:
                mail.send(msg)
                flash("Password reset instructions sent to your email")
                return redirect(url_for("forgot_password"))
            except Exception as e:
                flash(f"Error in sending email: {str(e)}")
                return redirect(url_for('forgot_password'))
        else:
            flash("User not found")
            return redirect(url_for('forgot_password'))

    return render_template("forgot_password.html")



from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet

@app.route('/generate_pdf', methods=['GET'])
@login_required
def generate_pdf():
    # Fetch answers and submissions for the current logged-in user
    answer_doc = mongo.db.answers.find_one({"user_id": current_user.id, "tender": "Proqurement"}) or {}
    submissions = answer_doc.get("submission_details", {})
    answers_dict = answer_doc.get("answers", {})
    selected_materials = answers_dict.get("selected_materials", [])

    # Create an in-memory PDF file
    pdf_stream = io.BytesIO()
    c = canvas.Canvas(pdf_stream, pagesize=letter)
    c.setFont("Helvetica", 12)
    y_position = 750  # Starting Y position

    # Title (reduced left indent to 50)
    c.drawString(50, y_position, 'Copy of Submitted options')
    y_position -= 30

    # Submission Details Section (reduced left indent to 50)
    if submissions:
        c.drawString(50, y_position, "Submission Details")
        y_position -= 20
        c.drawString(50, y_position, f"RFX Number: {submissions.get('rfx_number', 'N/A')}")
        y_position -= 20
        c.drawString(50, y_position, f"E-Tender Number: {submissions.get('etender_number', 'N/A')}")
        y_position -= 20
        c.drawString(50, y_position, f"Tender Call Number: {submissions.get('tendercall_number', 'N/A')}")
        y_position -= 30

    # Tender Section (reduced left indent to 50)
    tender = "Proqurement"
    c.drawString(50, y_position, f"Tender: {tender.capitalize()}")
    y_position -= 20

    # Regular answers (excluding selected_materials)
    for question, answer in answers_dict.items():
        if question != "selected_materials":
            c.drawString(50, y_position, f"{question}: {answer}")
            y_position -= 20

    # Selected Materials Section (Using Table for better formatting with adjusted left indent and wider columns)
    if selected_materials:
        c.drawString(50, y_position, "Selected items")
        y_position -= 20

        # Prepare table data
        table_data = [
            ['Material Name', 'Product ID', 'Quantity', 'Delivery Schedule']
        ]
        for item in selected_materials:
            delivery_schedule = "CM: [    ]Nos. within 2 months from date of LOA\nCP: @ [   ] Nos. per month thereafter"
            table_data.append([
                item.get('material', ''),
                item.get('product_id', ''),
                item.get('qty_advertised', ''),
                delivery_schedule
            ])

        # Create table with wider column widths and improved text wrapping
        col_widths = [250, 150, 100, 150]  # Further increased Material Name width to handle longer text
        table = Table(table_data, colWidths=col_widths)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),  # Reduced font size for better fitting
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('WORDWRAP', (0, 0), (-1, -1), True),  # Enable word wrapping
            ('LEFTPADDING', (0, 0), (-1, -1), 5),  # Add padding to prevent overlap
            ('RIGHTPADDING', (0, 0), (-1, -1), 5),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))

        # Calculate table height and position it (adjusted left indent to 50)
        # Use table's wrap to dynamically calculate height
        table.wrapOn(c, 50, y_position)  # Use dynamic wrap to determine height
        table_height = table._height  # Get the actual height after wrapping
        if y_position - table_height < 50:  # Check for page break
            c.showPage()
            c.setFont("Helvetica", 12)
            y_position = 750

        table.wrapOn(c, 50, y_position - table_height)  # Adjusted left indent to 50
        table.drawOn(c, 50, y_position - table_height)  # Adjusted left indent to 50
        y_position -= table_height + 20

    # Ensure we don't go off the page
    if y_position < 50:
        c.showPage()
        c.setFont("Helvetica", 12)
        y_position = 750

    # Save and finalize PDF
    c.showPage()
    c.save()
    pdf_stream.seek(0)

    return send_file(
        pdf_stream,
        as_attachment=True,
        download_name="review_answers.pdf",
        mimetype="application/pdf"
    )

# Reset password route
@app.route("/reset-password", methods=["GET", "POST"])
def reset_password():
    if request.method == "POST":
        new_password = request.form.get("new_password")
        username = request.form.get("username")

        # Check if password field is empty
        if not new_password:
            flash("Password cannot be empty")
            return redirect(url_for("reset_password"))

        hashed_password = generate_password_hash(new_password)

        account = mongo.db.accounts.find_one({"username": username})
        if account:
            mongo.db.accounts.update_one({"username": username}, {"$set": {"password": hashed_password}})
            flash("Your password has been changed")
            return redirect(url_for('login'))
        else:
            flash("User not found")
            return redirect(url_for('reset_password'))

    return render_template("reset_password.html")

from flask import send_file, redirect, url_for, flash
from flask_login import login_required, current_user
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO

@app.route("/annexure_f_docx")
@login_required
def annexure_f_docx():
    """
    Generates a DOCX file (Annexure-F format) with user data in a table,
    places e_tender_no / rfx_no from DB in the header (right side),
    and fills each row's Delivery Schedule with the CM/CP text.
    """

    # 1. Fetch user data from DB (Proqurement tender)
    doc_record = mongo.db.answers.find_one({"user_id": current_user.id, "tender": "Proqurement"})
    if not doc_record:
        flash("No data found for this tender.")
        return redirect(url_for("review"))

    # 2. Extract e_tender_no, rfx_no from the DB record
    e_tender_no = doc_record.get("etender_number", "SP/T-XXXX/XXXX")
    rfx_no      = doc_record.get("rfx_number",     "500000XXXX")

    # 3. Extract the selected materials from answers
    selected_materials = doc_record["answers"].get("selected_materials", [])

    # 4. Create a new Word document
    doc = Document()

    # Optional: set Normal style font
    style = doc.styles['Normal']
    style.font.name = 'Helvetica'
    style.font.size = Pt(11)

    # 5. Configure the header to display top-right text from DB
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Put them on two lines: e.g. "SP/T-0622/0724\nRFx No. 5000001271"
    run = header_para.add_run(f"{e_tender_no}\nRFx No. {rfx_no}")
    run.font.size = Pt(11)
    run.font.bold = True

    # 6. Body: Title "ANNEXURE-F" and Explanation
    doc.add_paragraph()  # blank paragraph for spacing

    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_title = title_para.add_run(
        "ANNEXURE-F\n\n"
        "Confirmation for Delivery/ Delivery Schedule\n"
        "[To be Uploaded and submitted with Techno-Commercial Bid] Refer Clause No.4.0 of section-III\n"
    )
    run_title.font.bold = True
    run_title.font.size = Pt(12)

    doc.add_paragraph()  # blank paragraph for spacing

    # 7. Create the table (6 columns)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Optional column widths
    table.columns[0].width = Inches(0.6)
    table.columns[1].width = Inches(1.0)
    table.columns[2].width = Inches(1.0)
    table.columns[3].width = Inches(1.1)
    table.columns[4].width = Inches(1.2)
    table.columns[5].width = Inches(1.9)  # Slightly bigger for multiline text

    # Fill header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Sr. No."
    hdr_cells[1].text = "Material"
    hdr_cells[2].text = "Product ID"
    hdr_cells[3].text = "Quantity Advertised (Nos.)"
    hdr_cells[4].text = "Quantity Offered by Bidders (Nos.)"
    hdr_cells[5].text = "Delivery Schedule"

    # 8. Add table rows from selected_materials
    for i, item in enumerate(selected_materials, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = item.get("material", "")
        row_cells[2].text = item.get("product_id", "")
        row_cells[3].text = item.get("qty_advertised", "")
        row_cells[4].text = item.get("qty_offered", "")

        # Insert multiline text for last column
        row_cells[5].text = (
            "CM:[    ]Nos. within 2 months \n"
            "from date of LOA \n"
            "CP: @ [   ] Nos. per \n"
            "month thereafter"
        )

    doc.add_paragraph()  # blank line

    # 9. Insert the note
    note_paragraph = doc.add_paragraph(
        "Note: The Entire Advertised Quantity of the tendered items to be supplied within 04 (Four) months from date of LOA. "
    )
    note_paragraph.add_run(
        "The above offered delivery against the subject tender should be as per terms and conditions of delivery as stated in Section III. "
    )
    note_paragraph.add_run(
        "In case of any disparity in the delivery schedule after opening of price bid, the above stated delivery shall be treated as final for evaluation of offer.\n"
    )

    doc.add_paragraph()  # blank line

    # 10. Signature lines (Left-aligned or right-aligned as needed)
    sig1 = doc.add_paragraph("Signature of the Tenderer: ________________")
    sig1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    sig2 = doc.add_paragraph("Designation: ____________________________")
    sig2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    sig3 = doc.add_paragraph("Date: ___________________")
    sig3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    sig4 = doc.add_paragraph("Seal of Company: ___________________")
    sig4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 11. Save to an in-memory BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # 12. Return the doc as a download
    return send_file(
        file_stream,
        as_attachment=True,
        download_name="ANNEXURE-F.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )



import os
from flask import send_file, redirect, url_for, flash, current_app
from flask_login import login_required, current_user
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING

def ordinal_call(num_str: str) -> str:
    """
    Convert an integer string like '5' into '5th',
    '3' -> '3rd', '2' -> '2nd', '1' -> '1st', etc.
    """
    try:
        call_num = int(num_str)
    except:
        call_num = 1  # fallback if parse fails

    if call_num == 1:
        return "1st"
    elif call_num == 2:
        return "2nd"
    elif call_num == 3:
        return "3rd"
    else:
        return f"{call_num}th"


@app.route("/annexure_g_docx")
@login_required
def annexure_g_docx():
    """
    Generates an ANNEXURE-G DOCX with:
      - A right-aligned header on every page, e.g.:
            SP/T-XXXX/XXXX [Nth Call]
            (RFX No.XXXXX)
      - Title: "ANNEXURE-G" (centered)
      - Subheading (centered)
      - Body paragraphs
      - Signature lines near bottom-right
    """

    # 1. Fetch the same doc_record from 'answers' for the "Proqurement" tender
    doc_record = mongo.db.answers.find_one({
        "user_id": current_user.id,
        "tender": "Proqurement"
    })
    if not doc_record:
        flash("No data found for this tender.")
        return redirect(url_for("review"))

    # 2. Extract needed fields from the DB record
    base_etender    = doc_record.get("etender_number",    "SP/T-XXXX/XXXX")
    rfx_no          = doc_record.get("rfx_number",        "500000XXXX")
    tendercall_str  = doc_record.get("tendercall_number", "1")
    # Convert call number to ordinal (e.g. "5" -> "5th") and build e_tender_no
    call_ordinal    = ordinal_call(tendercall_str)
    e_tender_no     = f"{base_etender} [{call_ordinal} Call]"

    # 3. Create a new Word Document
    doc = Document()

    # Adjust page margins
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

    # Base font style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # 4. Build a right-aligned header on every page:
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Insert the two lines in the header:
    # e.g. "SP/T-0622/0724 [5th Call]\n(RFX No.1002)"
    run_header = header_para.add_run(f"{e_tender_no}\n(RFX No.{rfx_no})")
    run_header.font.bold = True
    run_header.font.size = Pt(11)

    # 5. Title: "ANNEXURE-G" (centered, bold, 14pt)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_title = title_para.add_run("ANNEXURE-G")
    run_title.font.bold = True
    run_title.font.size = Pt(14)

    # 6. Subheading (centered, bold, 12pt)
    subheading_para = doc.add_paragraph()
    subheading_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = subheading_para.add_run(
        "Undertaking to be submitted by the Bidder declaring that Bidder is not "
        "Debarred/Blacklisted by Government/Semi-Government/Other Power Utilities"
    )
    sub_run.font.bold = True
    sub_run.font.size = Pt(12)

    # 7. Body paragraphs (single-spaced, minimal spacing)
    body_texts = [
        ("I/We hereby declare that I/We is/are participating in MSETCL’s "
         "Tender No. _______________."),
        ("As on date of submission of this Tender, I/We hereby declare that "
         "My Firm/We is/are not Debarred/Blacklisted by Any Government/ "
         "Semi-Government/Other Power Utilities, anywhere."),
        "The above declaration is true to the best of My/Our knowledge and belief.",
        ("I/We hereby agree that in case My Firm/We are Debarred/Blacklisted by "
         "Any Government/ Semi-Government/Other Power Utilities, anywhere, "
         "My/Our Offer is liable for rejection at any stage of Tendering process "
         "as per Tender Conditions."),
        ("Further, I/We hereby understand and agree that in case My Firm/We are "
         "Debarred/Blacklisted by Any Government/ Semi-Government/Other Power "
         "Utilities, anywhere, My/Our Order is liable for termination at any "
         "stage of Order execution process and My Firm/We shall be solely "
         "responsible for the consequences arising out of it."),
    ]

    for paragraph_text in body_texts:
        p = doc.add_paragraph(paragraph_text)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(6)

    # 8. Insert extra blank space before signature lines (adjust as needed)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(100)  # ~1.39 inch

    # 9. Signature lines near bottom-right
    sig_para1 = doc.add_paragraph()
    sig_para1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    sig_run1 = sig_para1.add_run("Authorized Signatory")
    sig_run1.font.bold = True
    sig_run1.font.size = Pt(12)

    sig_para2 = doc.add_paragraph()
    sig_para2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    sig_run2 = sig_para2.add_run("Seal of the Firm")
    sig_run2.font.bold = True
    sig_run2.font.size = Pt(12)

    # 10. Save to in-memory file
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # 11. Return as a download
    return send_file(
        file_stream,
        as_attachment=True,
        download_name="ANNEXURE-G.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


import os
from datetime import datetime
from flask import send_file, redirect, url_for, flash, current_app
from flask_login import login_required, current_user
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START

def ordinal_call(num_str: str) -> str:
    """
    Convert an integer string like '5' into '5th',
    '3' -> '3rd', '2' -> '2nd', '1' -> '1st', etc.
    """
    try:
        call_num = int(num_str)
    except:
        call_num = 1  # fallback if parse fails

    if call_num == 1:
        return "1st"
    elif call_num == 2:
        return "2nd"
    elif call_num == 3:
        return "3rd"
    else:
        return f"{call_num}th"

@app.route("/ANNEXURE_A")
@login_required
def ANNEXURE_A():
    """
    Generates a 2-page DOCX:
      - Right-aligned header on each page: E-Tender No. & RFX No.
      - Page 1: 'Schedule of E-Tender Notice' table + notes
      - Page 2: ANNEXURE ‘A’ Section III with special terms paragraphs.
      The 'tendercall_number' is appended to e_tender_no, e.g. "SP/T-0622/0724 [5th Call]".
    """

    # 1) Fetch DB record for "Proqurement" tender in 'answers' collection
    doc_record = mongo.db.answers.find_one({
        "user_id": current_user.id,
        "tender": "Proqurement"
    })
    if not doc_record:
        flash("No data found for this tender.")
        return redirect(url_for("review"))

    # 2) Extract the values from doc_record
    base_etender   = doc_record.get("etender_number",    "SP/T-XXXX/XXXX")
    rfx_no         = doc_record.get("rfx_number",        "500000XXXX")
    tendercall_str = doc_record.get("tendercall_number", "1")

    # Convert call number to ordinal, e.g. "5" => "5th", then build "SP/T-0622/0724 [5th Call]"
    call_ordinal = ordinal_call(tendercall_str)
    e_tender_no  = f"{base_etender} [{call_ordinal} Call]"

    # If you want the first selected material in the "Subject" text
    selected_materials = doc_record["answers"].get("selected_materials", [])
    if selected_materials:
        first_material = selected_materials[0].get("material", "Lightning Arrestors")
        subject_text = (
            f"Procurement of {first_material}\n"
            "against R&M scheme for various substations\n"
            "under EHV Zones in MSETCL"
        )
    else:
        subject_text = "Procurement of various Lightning Arrestors under R&M scheme"

    # Some example fields for page 1
    estimated_cost      = "Approx. Rs. 7.48 Crores"
    emd                 = "Rs. 50,000/-"
    tender_fee          = "Rs. 5,000/- + GST"
    due_date_submission = "08.08.2024, 17:00 Hrs"
    opening_date_time   = "08.08.2024, 17:45 Hrs"
    contact_person_info = (
        "E.E. Group (VI), CPA, C.O. MSETCL,\n"
        "1st Floor, Prakashgad Building, Bandra (E),\n"
        "Mumbai\n"
        "Email id: eegrp6@mahatransco.in"
    )

    # 3) Create the Word document
    doc = Document()

    # Adjust page margins
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(0.7)
    section.right_margin  = Inches(0.7)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # 4) Add a right-aligned header on each page
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Insert text: e.g. "SP/T-0622/0724 [5th Call]\n(RFX No.1002)"
    run_header = header_para.add_run(f"{e_tender_no}\n(RFX No.{rfx_no})")
    run_header.font.bold = True
    run_header.font.size = Pt(11)

    #
    # =========== PAGE 1 CONTENT ===========
    #

    # Insert MSETCL logo in the body
    logo_path = os.path.join(current_app.static_folder, "photos", "MSETCL.png")
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_logo = logo_para.add_run()
    run_logo.add_picture(logo_path, width=Inches(1.5))

    # Main heading (center)
    heading_para = doc.add_paragraph()
    heading_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_heading = heading_para.add_run("MAHARASHTRA STATE ELECTRICITY TRANSMISSION COMPANY LTD")
    run_heading.font.bold = True
    run_heading.font.size = Pt(14)

    # Subheading
    subheading_para = doc.add_paragraph()
    subheading_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_subheading = subheading_para.add_run("Schedule of E-Tender Notice")
    run_subheading.font.bold = True
    run_subheading.font.size = Pt(12)

    # Intro paragraph
    doc.add_paragraph(
        "The tender document is available on the website https://srmetender.mahasrtransco.in. "
        "Summary of the tender is given below."
    )

    # Table for e-tender details (9 rows)
    detail_table = doc.add_table(rows=1, cols=3)
    detail_table.style = 'Table Grid'
    hdr_cells = detail_table.rows[0].cells
    hdr_cells[0].text = "Sr."
    hdr_cells[1].text = "Tender Reference"
    hdr_cells[2].text = "Details"

    details_data = [
        ("1", "E-Tender No.",                         e_tender_no),
        ("2", "RFX No.",                              rfx_no),
        ("3", "Subject of E-Tender",                  subject_text),
        ("4", "Estimated Cost of Tender",             estimated_cost),
        ("5", "Tender Fee",                           tender_fee),
        ("6", "EMD",                                  emd),
        ("7", "Due Date & Time of Submission of Bid", due_date_submission),
        ("8", "Date & Time of Opening of Bid",        opening_date_time),
        ("9", "Contact Person",                       contact_person_info),
    ]
    for row_data in details_data:
        row_cells = detail_table.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]

    # Note paragraphs (a, b, c, d)
    note_para = doc.add_paragraph()
    note_para.add_run("Note:\n").bold = True
    note_para.add_run(
        "a. The tender documents can be downloaded only online from aforesaid website, "
        "within period of advertisement of tender.\n"
    )
    note_para.add_run(
        "b. The bidder shall submit their appropriate bids (Technical & Commercial) online, "
        "well in advance within due date & time indicated in the RFx. "
    )
    red_run = note_para.add_run("No extension to bid submission date & time will be given.\n")
    red_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    red_run.font.bold = True
    note_para.add_run(
        "c. The Techno-Commercial Bid and Price Bid will be opened on SRM e-Tendering portal.\n"
    )
    note_para.add_run(
        "d. MSETCL will not be responsible for non-submission of bid due to any website related issues."
    )

    #
    # =========== PAGE 2 CONTENT ===========
    #
    doc.add_page_break()

    # Page-2 heading
    second_page_title = doc.add_paragraph()
    second_page_title.paragraph_format.space_before = Pt(0)
    second_page_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_sec_title = second_page_title.add_run("ANNEXURE ‘A’\nSECTION III")
    run_sec_title.font.bold = True
    run_sec_title.font.size = Pt(14)

    spec_heading = doc.add_paragraph()
    spec_heading.paragraph_format.space_before = Pt(4)
    spec_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_spec = spec_heading.add_run("SPECIAL TERMS AND CONDITIONS AND INSTRUCTIONS TO THE TENDERERS")
    run_spec.font.bold = True
    run_spec.font.size = Pt(12)

    # Paragraph 1
    doc.add_paragraph(
        "The following Special Terms & Conditions shall form a part of the Tender Documents. "
        "The terms and conditions specified herein shall prevail in event of any conflict "
        "with existing Clauses as stated under Section ‘I’, ‘II’ and Technical Specification "
        "(Annexure ‘D’). The offer that do not adhere to stipulations of Section III shall "
        "be summarily rejected."
    )

    # Paragraph 2: 1.0 PRICE & PRICE VARIATION
    para_price = doc.add_paragraph()
    bold_price_title = para_price.add_run("1.0 PRICE & PRICE VARIATION:\n\n")
    bold_price_title.font.bold = True

    para_price.add_run(
        "The tenderers are requested to quote the prices on Variable Price Basis. The price "
        "variation shall be as per PV Formula indicated in Annexure ‘E’ of the tender document "
        "and the raw material price shall be as per monthly indices published by IEEMA. The "
        "same shall be capped at 20% on Positive Price Variation i.e. where the price trend is "
        "increasing. There shall be no lower ceiling/ capping on Negative Price Variation i.e. in "
        "case of savings where the price trend is decreasing.\n"
        "The Price Variation will be made applicable only if supply in all respects is completed "
        "by tenderer within the contractual period or within extended time period approved "
        "by the Competent Authority, if any. However, where the supplies are effected beyond the "
        "contractual delivery period or beyond the extended time period approved by the "
        "competent authority, if any, the purchaser shall remain entitled to savings on account of "
        "Negative Price Variation i.e. where the price trend is decreasing.\n"
        "The Price Variation shall be calculated from date of Placement of order (LOA) or Four "
        "(04) Months after Published Date of opening (Including Extension, if any) of techno- "
        "Commercial Bid of tender whichever is earlier.\n"
        "The prices quoted shall be on Free delivery by Road Transport to our various Stores/Sites "
        "anywhere in Maharashtra. The quantity offered and price shall be quoted at Item Level "
        "‘Conditions’ of the RFx.\n"
        "The offers in which the price bids are not submitted as stated above shall be liable for "
        "rejection."
    )

    # Paragraph 3: 2.0 EARNEST MONEY DEPOSIT
    para_emd = doc.add_paragraph()
    bold_emd_title = para_emd.add_run("2.0 EARNEST MONEY DEPOSIT:\n\n")
    bold_emd_title.font.bold = True

    para_emd.add_run(
        "The tenderers shall pay the Earnest Money Deposit @1% of estimated value of the "
        "tender as indicated in RFx advertised and the same shall be paid Online by the bidder "
        "before submission of their response on SRM e-tendering.\n"
        "For tenderers exempted from payment of EMD against their registration with "
        "MSME/SSI/NSIC, as per stipulations of Cl. No. 9.0 (b) of Section – I of tender "
        "document, the exemption shall only be applicable to the item(s)/ service(s) indicated in "
        "the registration certificate issued by the concerned authority. Note that, if the tendered "
        "item(s) is not maintained in the above list of item(s)/ service(s), the bidder shall be liable "
        "to pay the Earnest Money Deposit through “EMD Payment” link provided on SRM e- "
        "tendering before submission of their response. Failure to make payment of Earnest "
        "Money Deposit shall result in disqualification and rejection of offer thereof. The bidder "
        "shall submit the MSME/SSI/NSIC certificate alongwith their offer towards proof of "
        "exemption in payment of EMD.\n"
        "Other stipulations in this regard shall be as per Cl. No. 9.0 of Section –I (Annexure ‘A’)."
    )
    # Save doc to memory
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="SECTION-III.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )  

if __name__ == '__main__':
    app.run(debug=True, use_reloader=True, port=5000)
    